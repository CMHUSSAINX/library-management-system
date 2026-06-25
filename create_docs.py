import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_document():
    doc = docx.Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Color Palette (Corporate Navy and Teal)
    NAVY = RGBColor(15, 23, 42)
    TEAL = RGBColor(20, 184, 166)
    DARK_GRAY = RGBColor(51, 65, 85)
    
    # Set default paragraph style font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = DARK_GRAY

    # Helper function to add styled headings
    def add_heading_1(text):
        h = doc.add_paragraph()
        run = h.add_run(text)
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = NAVY
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        return h

    def add_heading_2(text):
        h = doc.add_paragraph()
        run = h.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = TEAL
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True
        return h

    def add_heading_3(text):
        h = doc.add_paragraph()
        run = h.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = NAVY
        h.paragraph_format.space_before = Pt(8)
        h.paragraph_format.space_after = Pt(2)
        h.paragraph_format.keep_with_next = True
        return h

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("Library Management System")
    title_run.bold = True
    title_run.font.size = Pt(26)
    title_run.font.color.rgb = NAVY
    title_p.paragraph_format.space_after = Pt(2)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run("SOLID Full-Stack System Specifications & User Roles Guide")
    sub_run.italic = True
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = TEAL
    subtitle_p.paragraph_format.space_after = Pt(24)

    # Document Divider Line
    divider = doc.add_paragraph()
    divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
    divider.add_run("__________________________________________________________________").font.color.rgb = TEAL
    divider.paragraph_format.space_after = Pt(24)

    # Overview Section
    add_heading_1("1. System Overview")
    doc.add_paragraph(
        "This Library Management System is a robust, full-stack web application designed for modern library cataloging, "
        "book tracking, and member borrowings. The system is engineered to satisfy high-performance standards, "
        "strict data validation, and clean interface usability for both administrators and regular readers."
    )
    
    # Tech Stack Table
    add_heading_2("Technology Stack Details")
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Light Shading Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Layer'
    hdr_cells[1].text = 'Technology Used'
    hdr_cells[2].text = 'Key Details / Packages'
    
    # Style Table Header
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = NAVY
        
    row_data = [
        ('Frontend', 'React.js (Vite)', 'React Router DOM, Axios (API Clients), Lucide React (Icons), Vanilla CSS Modules'),
        ('Backend', 'Node.js + Express', 'jsonwebtoken (JWT), bcryptjs (Password Hashing), dotenv, cors'),
        ('Database', 'MongoDB', 'Mongoose ODM, Hosted on MongoDB Atlas Cloud Cluster')
    ]
    
    for i, data in enumerate(row_data):
        row_cells = table.rows[i+1].cells
        row_cells[0].text = data[0]
        row_cells[1].text = data[1]
        row_cells[2].text = data[2]
        
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # SOLID Principles Section
    add_heading_1("2. Application of SOLID Design Principles")
    doc.add_paragraph(
        "To maximize codebase clean-cut separation, lower logical coupling, and optimize growth support, "
        "the backend is strictly architected around the five SOLID software design principles:"
    )

    add_heading_3("Single Responsibility Principle (SRP)")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Mongoose Models (User, Book, BorrowRecord) ").bold = True
    p.add_run("are strictly dumb definitions of database schemas and data shapes.")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Service Layers (AuthService, BookService, BorrowService) ").bold = True
    p.add_run("house 100% of the library business logic, calculations, and rules.")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Controllers (authController, bookController, borrowController) ").bold = True
    p.add_run("solely handle incoming HTTP requests, validate input format, and return status codes/JSON responses.")

    add_heading_3("Open/Closed Principle (OCP)")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Express Middleware Chain: ").bold = True
    p.add_run("Routes are configured to accept pluggable middleware modules (e.g. JWT filters, role restriction checks) "
              "which allows extending route execution pipelines without modifying the controller classes themselves.")

    add_heading_3("Liskov Substitution Principle (LSP)")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Consistent Contracts: ").bold = True
    p.add_run("All service methods are structured to return predictable, strongly shaped Data Transfer Objects (DTOs) "
              "or Mongoose documents. Any implementation of these services can be substituted without breaking the client-facing APIs.")

    add_heading_3("Interface Segregation Principle (ISP)")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Granular Services: ").bold = True
    p.add_run("Instead of a single heavy library service, functionalities are strictly segregated into domain-focused services: "
              "Auth (credentials), Book (cataloging and search), and Borrow (lending transactions). Route directories map "
              "specifically to these endpoints so client interfaces do not depend on methods they do not call.")

    add_heading_3("Dependency Inversion Principle (DIP)")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Service Inversion: ").bold = True
    p.add_run("High-level API endpoints (Controllers) depend on high-level business abstractions (Services) rather than "
              "directly querying Mongoose queries in routes. Additionally, database connection setup configures a DNS-server resolution "
              "override to guarantee connections to cloud clusters can be handshaked without dependency on local router configurations.")

    # Functionalities Section
    add_heading_1("3. System Functionalities & User Roles")
    doc.add_paragraph(
        "The system separates user actions based on roles (Admin vs. Member), authenticated via JSON Web Tokens (JWT) "
        "and secured by Express middleware filters."
    )

    add_heading_2("A. Administrator Role (Admin Panel)")
    doc.add_paragraph("Administrators manage the inventory and track overall lending activities. Key actions include:")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Book Cataloging (Add Book): ").bold = True
    p.add_run("Register new titles with title, author, unique ISBN, genre description, and total stock copies.")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Edit Catalog details: ").bold = True
    p.add_run("Update authors, titles, or adjust total copy count. Adjusting copy count automatically manages available "
              "copies safely, preventing reduction below the number of currently checked-out copies.")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Delete Book: ").bold = True
    p.add_run("Remove volumes from circulation. The system prevents deletions if any reader has currently borrowed copies of that book.")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("View Active Borrow Logs: ").bold = True
    p.add_run("Monitor all active transactions in the library, displaying reader name, email, book title, check-out date, "
              "and return due date (automatically highlighting overdue items).")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("System Dashboard Metrics: ").bold = True
    p.add_run("Access real-time indicators for cataloged titles, available shelf stock, and total books out.")

    add_heading_2("B. Library Member Role (Reader Hub)")
    doc.add_paragraph("Members borrow books and track their personal check-outs. Key actions include:")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Account Sign Up & Login: ").bold = True
    p.add_run("Create secure accounts and log in to obtain session tokens, which are persisted locally.")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Search & Filter Catalog: ").bold = True
    p.add_run("Instantly search books by title, author name, or ISBN, and filter dynamically by genre categories.")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Borrow Book: ").bold = True
    p.add_run("Borrow a copy of an available book. The system decrements the available copy count on the shelf and "
              "creates an active borrow transaction. The system prevents borrowing a book the reader currently has checked out.")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Return Book: ").bold = True
    p.add_run("Return the book from the 'My Borrowings' dashboard. The system records the return date, updates status "
              "to returned, and increments available shelf copies.")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Track Return Due Dates: ").bold = True
    p.add_run("View active check-outs. Borrowings have a standard 14-day return period. Overdue books are highlighted "
              "in red alerts on the user's dashboard.")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Personal Borrowing History: ").bold = True
    p.add_run("Access a chronological list of all past library check-outs and returned dates.")

    # Save document
    doc.save("Library_Management_System_Documentation.docx")
    print("Documentation word file created successfully!")

if __name__ == "__main__":
    create_document()
